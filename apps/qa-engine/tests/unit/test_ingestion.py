"""Unit tests for document parsing and requirement persistence.

Covers tools.file_ingestion (FR-IN-001..004, SEC-011) and
app.services.ingestion (FR-IN-005 duplicate detection, FR-IN-006
versioning) against an in-memory SQLite database — no network, no LLM
(SRS §15.1, NFR-MNT-003).
"""

from __future__ import annotations

import io
import json

import pytest
from sqlalchemy import create_engine, select
from sqlalchemy.orm import sessionmaker

from app.models.db import Base
from app.models.entities import Project, Requirement
from app.services.ingestion import (
    DuplicateRequirementError,
    create_requirement_from_text,
    ingest_file,
)
from tools.file_ingestion import (
    MAX_FILE_SIZE_BYTES,
    IngestionError,
    parse_document,
)

pytestmark = pytest.mark.unit


# ---------------------------------------------------------------------------
# Parsing (tools.file_ingestion)
# ---------------------------------------------------------------------------


class TestParseTxt:
    def test_paragraph_chunks_and_traceability(self, tmp_path):
        f = tmp_path / "reqs.txt"
        f.write_text(
            "Users must be able to log in.\n\nUsers must be able to log out.\n",
            encoding="utf-8",
        )
        result = parse_document(f)
        assert result["kind"] == "text"
        assert result["source"] == "reqs.txt"
        assert "log in" in result["text"]
        assert len(result["chunks"]) == 2
        # FR-IN-003: every chunk records source + location.
        assert result["chunks"][0]["location"] == "paragraph 1"
        assert all(c["source"] == "reqs.txt" for c in result["chunks"])
        assert result["metadata"]["paragraph_count"] == 2

    def test_metadata_has_size_and_hash(self, tmp_path):
        f = tmp_path / "one.txt"
        f.write_text("Only one requirement here.", encoding="utf-8")
        meta = parse_document(f)["metadata"]
        assert meta["size_bytes"] == f.stat().st_size
        assert len(meta["sha256"]) == 64


class TestParseMarkdown:
    def test_one_chunk_per_top_level_heading(self, tmp_path):
        f = tmp_path / "spec.md"
        f.write_text(
            "# Login\nUsers must log in.\n\n## Details\nEmail plus password.\n\n"
            "# Logout\nUsers can log out.\n",
            encoding="utf-8",
        )
        result = parse_document(f)
        assert result["kind"] == "markdown"
        headings = [c["heading"] for c in result["chunks"]]
        assert headings == ["Login", "Logout"]
        # The sub-heading stays inside the first section body.
        assert "Details" in result["chunks"][0]["text"]

    def test_markdown_without_headings_falls_back_to_body_chunk(self, tmp_path):
        f = tmp_path / "notes.md"
        f.write_text("Just some requirement prose.\n\nAnd a second paragraph.\n", encoding="utf-8")
        result = parse_document(f)
        assert result["metadata"]["heading_count"] == 0
        # Without headings the whole body becomes one heading-less chunk.
        assert len(result["chunks"]) == 1
        assert result["chunks"][0]["heading"] == ""
        assert "second paragraph" in result["chunks"][0]["text"]


class TestParseJson:
    def test_story_list_normalised(self, tmp_path):
        stories = [
            {
                "title": "Login",
                "text": "As a user I want to log in.",
                "acceptance_criteria": ["Valid creds open the item list"],
            },
            {"story": "As a user I want to reset my password."},
        ]
        f = tmp_path / "stories.json"
        f.write_text(json.dumps(stories), encoding="utf-8")
        result = parse_document(f)
        assert result["kind"] == "json"
        parsed = result["metadata"]["stories"]
        assert len(parsed) == 2
        assert parsed[0]["title"] == "Login"
        assert parsed[0]["acceptance_criteria"] == ["Valid creds open the item list"]
        # Untitled story derives its title from the text (FR-IN-002).
        assert parsed[1]["title"].startswith("As a user I want to reset")

    def test_invalid_json_rejected_with_actionable_message(self, tmp_path):
        f = tmp_path / "broken.json"
        f.write_text("{not json", encoding="utf-8")
        with pytest.raises(IngestionError, match="not valid JSON"):
            parse_document(f)

    def test_story_without_text_rejected(self):
        data = json.dumps([{"title": "Empty story"}]).encode()
        with pytest.raises(IngestionError, match="no 'text' or 'story'"):
            parse_document(data, "stories.json")


class TestParseDocx:
    def test_docx_built_on_the_fly(self):
        import docx

        document = docx.Document()
        document.add_heading("Checkout", level=1)
        document.add_paragraph("Users can pay by card.")
        document.add_heading("Refunds", level=1)
        document.add_paragraph("Refunds are honoured within 30 days.")
        buffer = io.BytesIO()
        document.save(buffer)

        result = parse_document(buffer.getvalue(), "spec.docx")
        assert result["kind"] == "docx"
        headings = [c["heading"] for c in result["chunks"]]
        assert headings == ["Checkout", "Refunds"]
        assert "pay by card" in result["chunks"][0]["text"]
        assert all(c["source"] == "spec.docx" for c in result["chunks"])


class TestUploadValidation:
    def test_oversize_file_rejected(self):
        data = b"a" * (MAX_FILE_SIZE_BYTES + 1)
        with pytest.raises(IngestionError, match="exceeds the 10 MB limit"):
            parse_document(data, "big.txt")

    def test_empty_file_rejected(self):
        with pytest.raises(IngestionError, match="is empty"):
            parse_document(b"   \n  ", "empty.txt")

    def test_unsupported_extension_rejected(self):
        with pytest.raises(IngestionError, match="not supported"):
            parse_document(b"MZ...", "malware.exe")

    def test_bytes_without_filename_rejected(self):
        with pytest.raises(IngestionError, match="file name"):
            parse_document(b"some text")

    def test_missing_path_rejected(self, tmp_path):
        with pytest.raises(IngestionError, match="not found"):
            parse_document(tmp_path / "nope.txt")


# ---------------------------------------------------------------------------
# Persistence (app.services.ingestion) on in-memory SQLite
# ---------------------------------------------------------------------------


@pytest.fixture()
def db_session():
    """Isolated in-memory database session (SRS §15.1: offline unit tests)."""
    engine = create_engine("sqlite://")
    Base.metadata.create_all(engine)
    factory = sessionmaker(bind=engine, expire_on_commit=False)
    session = factory()
    try:
        yield session
    finally:
        session.close()
        engine.dispose()


@pytest.fixture()
def project(db_session):
    proj = Project(name="unit-test-project")
    db_session.add(proj)
    db_session.commit()
    return proj


class TestCreateRequirement:
    def test_first_version_is_one(self, db_session, project):
        req = create_requirement_from_text(
            db_session, project.id, "Login", "Users must be able to log in.", ["Valid creds work"]
        )
        assert req.version == 1
        assert req.source == "manual"
        assert req.acceptance_criteria == ["Valid creds work"]

    def test_exact_duplicate_rejected(self, db_session, project):
        create_requirement_from_text(db_session, project.id, "Login", "Users must log in.")
        # FR-IN-005: identical text in the same project is a duplicate even
        # under a different title.
        with pytest.raises(DuplicateRequirementError):
            create_requirement_from_text(
                db_session, project.id, "Login (copy)", "Users must log in."
            )

    def test_version_increment_keeps_old_row(self, db_session, project):
        v1 = create_requirement_from_text(db_session, project.id, "Login", "Users must log in.")
        v2 = create_requirement_from_text(
            db_session, project.id, "Login", "Users must log in with MFA."
        )
        # FR-IN-006: same title + source, new text -> new version, old kept.
        assert (v1.version, v2.version) == (1, 2)
        assert v1.id != v2.id
        rows = db_session.scalars(select(Requirement)).all()
        assert len(rows) == 2

    def test_empty_text_rejected(self, db_session, project):
        with pytest.raises(IngestionError, match="empty"):
            create_requirement_from_text(db_session, project.id, "Login", "   ")

    def test_unknown_project_rejected(self, db_session):
        with pytest.raises(IngestionError, match="not found"):
            create_requirement_from_text(db_session, "no-such-id", "Login", "Some text.")


class TestIngestFile:
    MD_V1 = b"# Checkout\nUsers can pay by card.\n\n# Refunds\nRefunds within 30 days.\n"
    MD_V2 = b"# Checkout\nUsers can pay by card or wallet.\n\n# Refunds\nRefunds within 30 days.\n"

    def test_one_requirement_per_heading(self, db_session, project):
        created = ingest_file(db_session, project.id, "spec.md", self.MD_V1)
        assert [r.title for r in created] == ["Checkout", "Refunds"]
        assert all(r.source == "spec.md" for r in created)
        assert all(r.version == 1 for r in created)

    def test_reingest_identical_file_rejected_as_duplicate(self, db_session, project):
        ingest_file(db_session, project.id, "spec.md", self.MD_V1)
        with pytest.raises(DuplicateRequirementError, match="already exist"):
            ingest_file(db_session, project.id, "spec.md", self.MD_V1)

    def test_reingest_changed_section_creates_new_version(self, db_session, project):
        ingest_file(db_session, project.id, "spec.md", self.MD_V1)
        created = ingest_file(db_session, project.id, "spec.md", self.MD_V2)
        # Only the changed 'Checkout' section is stored, as version 2;
        # the unchanged 'Refunds' section is skipped as a duplicate.
        assert [r.title for r in created] == ["Checkout"]
        assert created[0].version == 2
        checkouts = db_session.scalars(
            select(Requirement).where(Requirement.title == "Checkout")
        ).all()
        assert sorted(r.version for r in checkouts) == [1, 2]

    def test_json_stories_become_requirements(self, db_session, project):
        stories = json.dumps(
            [
                {
                    "title": "Search",
                    "text": "Users can search products.",
                    "acceptance_criteria": ["Results are ranked"],
                }
            ]
        ).encode()
        (req,) = ingest_file(db_session, project.id, "stories.json", stories)
        assert req.title == "Search"
        assert req.acceptance_criteria == ["Results are ranked"]
        assert req.source == "stories.json"
